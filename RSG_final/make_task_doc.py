from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

sec = doc.sections[0]
sec.page_width = Cm(21.0)
sec.page_height = Cm(29.7)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)
sec.top_margin = Cm(2.5)
sec.bottom_margin = Cm(2.5)

FONT = '맑은 고딕'


def _shading(p, fill):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    pPr.append(shd)


def run(p, text, size=10, bold=False, italic=False, color=None):
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = color
    return r


def para(text='', size=10, bold=False, italic=False, color=None,
         align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=4,
         left_indent=0, first_indent=0):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.3
    if left_indent:
        p.paragraph_format.left_indent = Inches(left_indent)
    if first_indent:
        p.paragraph_format.first_line_indent = Inches(first_indent)
    if text:
        run(p, text, size=size, bold=bold, italic=italic, color=color)
    return p


def section_title(text, fill='1F3D6E'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    _shading(p, fill)
    r = p.add_run(f'  {text}')
    r.font.name = FONT
    r.font.size = Pt(11)
    r.bold = True
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    return p


def person_block(name, role_text, fill):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    _shading(p, fill)
    r1 = p.add_run(f'  {name}')
    r1.font.name = FONT
    r1.font.size = Pt(11)
    r1.bold = True
    r1.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    r2 = p.add_run(f'  |  {role_text}')
    r2.font.name = FONT
    r2.font.size = Pt(9.5)
    r2.font.color.rgb = RGBColor(0xFF, 0xFF, 0xCC)
    return p


def bullet(text, size=10, indent=0.2, color=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.line_spacing = 1.3
    run(p, text, size=size, color=color)
    return p


def sub_bullet(text, size=9.5):
    p = doc.add_paragraph(style='List Bullet 2')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.line_spacing = 1.3
    run(p, text, size=size, color=RGBColor(0x33, 0x33, 0x33))
    return p


def callout(text, fill, text_color=RGBColor(0x11, 0x11, 0x11), size=9.5):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.35
    _shading(p, fill)
    run(p, text, size=size, color=text_color)
    return p


def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run(p, '─' * 60, size=8, color=RGBColor(0xBB, 0xBB, 0xBB))
    return p


# ══════════════════════════════════════════════════════
# HEADER
# ══════════════════════════════════════════════════════
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_after = Pt(4)
_shading(p_title, '1F3D6E')
r = p_title.add_run('  논문 공동 작업 업무 지시서  ')
r.font.name = FONT
r.font.size = Pt(14)
r.bold = True
r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_sub.paragraph_format.space_after = Pt(16)
_shading(p_sub, 'D9EAF7')
run(p_sub,
    '서울시 저소득 독거노인의 소득 기반 거주지 분리 측도 분析  |  작성일: 2026-06-09',
    size=9.5, color=RGBColor(0x1F, 0x3D, 0x6E))

# ══════════════════════════════════════════════════════
# 1. 제출 일정
# ══════════════════════════════════════════════════════
section_title('1. 제출 일정')

tbl = doc.add_table(rows=5, cols=3)
tbl.style = 'Table Grid'
headers = ['담당자', '마감일', '전달 대상']
for i, h in enumerate(headers):
    c = tbl.rows[0].cells[i]
    _shading(c.paragraphs[0], '2E5090')
    r = c.paragraphs[0].add_run(h)
    r.font.name = FONT
    r.font.size = Pt(10)
    r.bold = True
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

rows_data = [
    ('영서 · 유나', '6월 12일 (목) 자정', '서기'),
    ('승원',        '6월 12일 (목) 자정', '서기'),
    ('서기',        '6월 14일 (토) 오전 중', '진'),
    ('진 (최종 투고)', '6월 14일 (토) 마감', '학회지 시스템'),
]
fills = ['F5F8FF', 'F5F8FF', 'FFF9E6', 'FFF0F0']
for ri, (a, b, c_) in enumerate(rows_data):
    row = tbl.rows[ri + 1]
    for ci, val in enumerate([a, b, c_]):
        cell = row.cells[ci]
        _shading(cell.paragraphs[0], fills[ri])
        r = cell.paragraphs[0].add_run(val)
        r.font.name = FONT
        r.font.size = Pt(9.5)
        if ri == 3:
            r.bold = True

doc.add_paragraph().paragraph_format.space_after = Pt(4)

callout(
    '⚠️  마감일 엄수 요청. 서기의 최종 포맷팅 작업(목~토)을 위해 목요일 자정까지 반드시 전달 바랍니다.',
    'FFF3CD', RGBColor(0x7B, 0x4A, 0x00)
)

# ══════════════════════════════════════════════════════
# 2. 담당자별 업무 범위
# ══════════════════════════════════════════════════════
section_title('2. 담당자별 업무 범위')

# ── 영서 + 유나 ──
person_block('영서 · 유나', '서론 / 연구방법 / 결과 및 논의', '2E5090')
para('담당 섹션: 1. 서론  |  3. 연구방법  |  4. 결과 및 논의',
     size=9.5, bold=True, color=RGBColor(0x2E, 0x50, 0x90), space_before=4)

bullet('초안 가이드(논문_초안_v2.docx) 기반으로 각 단락 작성')
sub_bullet('각 단락 앞의 핵심 논리 / 확장 예시를 참고해 내용 방향 확인')
sub_bullet('확장 예시를 그대로 복붙하지 말고, 본인이 소화한 언어로 재작성')

bullet('인용 기반 서술은 반드시 원문 직접 확인 후 작성 — 이것이 가장 중요한 요청사항')
sub_bullet('[SW] 태그가 붙은 인용(승원씨 정리본 기반)은 원문 논문을 열어 해당 주장이 실제로 있는지 확인')
sub_bullet('확인된 인용만 남기고, 원문과 다르면 문장 자체를 수정')
sub_bullet('LLM이 인용 내용을 만들어냈을 가능성이 있으므로 출처 없는 주장은 삭제')

bullet('Figure 삽입 시 통일 기준 준수 (3절 참고)')
bullet('지도학회지 본문 작성 양식 준수 (4절 참고)')
bullet('인용 시 처음부터 지도학회지 편집규정 형식으로 직접 작성 — 본문 인용 달면서 동시에 맨 뒤 참고문헌 목록에도 추가')

# ── 승원 ──
person_block('승원', '참고문헌 정리 및 검증', '2E6B3A')
para('담당 섹션: 참고문헌 전체',
     size=9.5, bold=True, color=RGBColor(0x2E, 0x6B, 0x3A), space_before=4)

bullet('현재 초안에 사용된 모든 참고문헌의 서지 정보 최종 확인')
sub_bullet('저자명 / 연도 / 제목 / 저널명 / 권호 / 페이지 / DOI 완전히 맞는지 체크')
sub_bullet('⚠️ 팀원 확인 필요 항목(논문_초안_가이드.md 하단 테이블 참고) 우선 처리')

bullet('[SW] 태그 인용 7건 원문 확인 — 해외선행연구.xlsx에 정리해두었던 논문들')
sub_bullet('Cartone(2025), de Sousa Filho(2022), Owens(2016)')
sub_bullet('Qin et al.(2024), Menec(2019), Ku(2021), Kim et al.(2021)')
sub_bullet('각 논문 원문 확인 후: 맞으면 [SW] 태그 제거, 틀리면 수정 또는 삭제')

bullet('국내 문헌 서지 최종 확인')
sub_bullet('이희연·이다예·유재성(2015) — 정확한 제목·권호 확인')
sub_bullet('안용한·김영호(2023) — 정확한 제목·권호 확인')

bullet('지도학회지 편집규정 참고문헌 형식으로 직접 작성')
sub_bullet('국내: 저자, 연도, 제목, 저널명, 권(호), 페이지.')
sub_bullet('외국어: 저자, 연도, 제목, 저널명(이탤릭), 권(호), 페이지. DOI.')
sub_bullet('배열 순서: 국내 문헌 → 외국어 문헌 / 각각 저자명 자모음·알파벳 순')
bullet('정리된 참고문헌 목록 Word 파일로 서기에게 목요일까지 전달')

# ── 서기 ──
person_block('서기', '초록 · 결론 작성  +  전체 최종 포맷팅', '7B2D8B')
para('담당 섹션: 초록  |  결론  |  전체 최종 포맷 검토',
     size=9.5, bold=True, color=RGBColor(0x7B, 0x2D, 0x8B), space_before=4)

bullet('초록 및 결론 작성 (논문_초안_v2.docx 핵심 논리·확장 예시 참고)')
sub_bullet('초록: 4단락 구조 (문제 제기 → 방법론 → 핵심 결과 → 연구 의미)')
sub_bullet('결론: 연구 요약 → 이론적 기여 → 한계 → 향후 연구 순서')
sub_bullet('초록·결론도 수치 인용 포함되므로 본문과 수치 일치 여부 확인')

bullet('영서·유나·승원 전달 파일 취합 후 전체 포맷 최종 점검')
sub_bullet('지도학회지 양식 기준 폰트·여백·줄간격·제목 스타일 통일')
sub_bullet('Figure·Table 번호 순서, 캡션 형식 일치 여부 확인')
sub_bullet('참고문헌 목록과 본문 인용 표기 일치 여부 최종 대조 (누락·중복 확인)')
sub_bullet('⚠️ 표시 미해결 항목 없는지 전체 검색 후 처리')

bullet('완성 파일 토요일 오전 중 진에게 전달')

# ══════════════════════════════════════════════════════
# 3. 공통 주의사항
# ══════════════════════════════════════════════════════
section_title('3. 공통 주의사항')

# 3-1 팩트체크
para('① 인용 팩트체크 — 가장 중요', size=10.5, bold=True,
     color=RGBColor(0xCC, 0x00, 0x00), space_before=8)
callout(
    '이 논문은 여러 명이 LLM을 활용해 협업한 결과물입니다. LLM 할루시네이션으로 인해 '
    '실제 논문에 없는 주장이 인용으로 달렸을 가능성이 있습니다. '
    '우리 이름으로 투고하는 논문이므로 인용된 모든 문헌은 원문을 직접 확인하고 작성해 주세요.',
    'FFF0F0', RGBColor(0x7B, 0x00, 0x00)
)
bullet('원문을 읽지 않고 인용을 달지 말 것')
bullet('확인 안 된 인용은 차라리 삭제하는 것이 나음')
bullet('문장이 맞더라도 인용 논문이 실제로 그 주장을 하는지 확인 필요')

# 3-2 양식
para('② 지도학회지 양식 준수', size=10.5, bold=True,
     color=RGBColor(0x1F, 0x3D, 0x6E), space_before=10)
bullet('투고 대상: 한국지도학회지 (KCI 등재지)')
bullet('양식 파일은 학회 홈페이지에서 다운로드 후 적용')
bullet('폰트: 본문 신명조 10pt / 영문 Times New Roman 10pt (학회 양식 확인 후 적용)')
bullet('줄간격, 여백, 제목 스타일 학회 양식 기준으로 통일')
bullet('서기가 최종 포맷팅 점검하지만, 작성 시부터 양식 맞춰서 주면 수고가 줄어듦')

# 3-3 Figure/Table
para('③ Figure · Table 통일 기준', size=10.5, bold=True,
     color=RGBColor(0x1F, 0x3D, 0x6E), space_before=10)
bullet('해상도: 300 dpi 이상 (지도 이미지 포함)')
bullet('폰트: 맑은 고딕 또는 나눔고딕 통일 (그림 안 텍스트 포함)')
bullet('범례 디자인: 색상 팔레트 통일 — 집중 극(붉은 계열) / 배제 극(파란 계열) / 중립(회색)')
bullet('캡션 형식: "Figure N. 내용 설명." / "Table N. 내용 설명." — 마침표로 끝냄')
bullet('그림 크기: 단단(single column) 또는 전단(full width) 중 하나로 통일')

# 3-4 LLM
para('④ LLM 활용 가이드', size=10.5, bold=True,
     color=RGBColor(0x1F, 0x3D, 0x6E), space_before=10)
bullet('LLM 활용 적극 권장 — 단, 결과물은 반드시 본인이 검토하고 수정')
bullet('AI 티 나는 표현 가급적 줄이기')
sub_bullet('"~임을 알 수 있다", "~라 할 수 있다", "~는 매우 중요하다" 등 완곡·과장 표현 자제')
sub_bullet('수동태 과다 사용 자제 — 능동형으로 직접 서술')
sub_bullet('한 문장이 지나치게 길어지지 않도록 (2~3행 이내)')
bullet('원고 제출 전 본인이 소리 내어 읽어보기 — 어색한 부분 자연스럽게 수정')

# 3-5 EndNote
para('⑤ 참고문헌 — 지도학회지 편집규정 형식으로 직접 작성', size=10.5, bold=True,
     color=RGBColor(0x1F, 0x3D, 0x6E), space_before=10)
callout(
    'EndNote 사용 없이 지도학회지 편집규정(제13조)에 따라 직접 작성합니다. '
    '본문에 인용 표기하면서 동시에 맨 뒤 참고문헌 단락에 해당 문헌을 추가해 주세요.',
    'EBF5EB', RGBColor(0x1A, 0x40, 0x20)
)
bullet('본문 인용 형식')
sub_bullet('서술형: 저자(연도) → 예: Feitosa(2007)는…')
sub_bullet('괄호형: (저자, 연도) → 예: …로 나타났다(Bell, 1954).')
sub_bullet('2인 저자: 외국어 and, 한국어 · → Reardon and Bischoff(2011) / 안용한·김영호(2023)')
sub_bullet('3인 이상: 외국어 et al., 한국어 등 → Kim et al.(2021) / 이희연 등(2015)')
sub_bullet('여러 문헌 괄호 안: 세미콜론(;)으로 분리, 국문 먼저 → (이희연 등, 2015; Bell, 1954)')
bullet('참고문헌 목록 형식 (맨 뒤에 직접 추가)')
sub_bullet('국내: 안용한·김영호, 2023, "제목," 저널명, 권(호), 시작-끝쪽.')
sub_bullet('외국어: Duncan, O.D. and Duncan, B., 1955, Title, Journal Name(이탤릭), 권(호), 쪽.')
sub_bullet('배열: 국내 문헌 먼저(자모음 순) → 외국어 문헌(알파벳 순)')
bullet('인용하면서 바로 목록에 추가하는 습관 — 나중에 모아서 정리하면 누락 발생')

# ══════════════════════════════════════════════════════
# 4. 전체 일정 요약
# ══════════════════════════════════════════════════════
section_title('4. 전체 일정 한눈에 보기')

tbl2 = doc.add_table(rows=7, cols=4)
tbl2.style = 'Table Grid'
for i, h in enumerate(['날짜', '담당', '할 일', '전달']):
    c = tbl2.rows[0].cells[i]
    _shading(c.paragraphs[0], '2E5090')
    r = c.paragraphs[0].add_run(h)
    r.font.name = FONT
    r.font.size = Pt(9.5)
    r.bold = True
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

sched = [
    ('6/9 (화) 오늘', '전원', '이 지시서 확인 및 초안 v2 파일 수령', '—'),
    ('6/9~6/12', '영서·유나', '서론·방법·결과 작성 + 인용 팩트체크', '→ 서기'),
    ('6/9~6/12', '승원', '참고문헌 전체 검증 + EndNote 라이브러리', '→ 서기'),
    ('6/12 (목) 자정', '영서·유나·승원', '마감 — 서기에게 파일 전달', '→ 서기'),
    ('6/12~6/14', '서기', '초록·결론 작성 + 전체 포맷팅 최종 점검', '→ 진'),
    ('6/14 (토) 오전', '서기 → 진', '최종 완성본 전달 및 투고 준비', '→ 학회'),
]
fills2 = ['F5F8FF', 'F5F8FF', 'F5F8FF', 'FFF3CD', 'FFF9E6', 'FFF0F0']
for ri, row_data in enumerate(sched):
    for ci, val in enumerate(row_data):
        cell = tbl2.rows[ri + 1].cells[ci]
        _shading(cell.paragraphs[0], fills2[ri])
        r = cell.paragraphs[0].add_run(val)
        r.font.name = FONT
        r.font.size = Pt(9)
        if ri in (3, 5):
            r.bold = True

doc.add_paragraph().paragraph_format.space_after = Pt(6)

# ══════════════════════════════════════════════════════
# 5. 전달 파일 안내
# ══════════════════════════════════════════════════════
section_title('5. 참고 파일 안내')

para('아래 파일들을 공유 드립니다. 작업 전 반드시 확인해 주세요.',
     size=9.5, space_before=4)

tbl3 = doc.add_table(rows=4, cols=3)
tbl3.style = 'Table Grid'
for i, h in enumerate(['파일명', '내용', '주요 사용자']):
    c = tbl3.rows[0].cells[i]
    _shading(c.paragraphs[0], '2E5090')
    r = c.paragraphs[0].add_run(h)
    r.font.name = FONT
    r.font.size = Pt(9.5)
    r.bold = True
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

files = [
    ('논문_초안_v2.docx',
     '전체 초안 + 단락별 핵심 논리·확장 예시 포함 가이드',
     '전원'),
    ('논문_초안_가이드.md',
     '섹션별 작성 원칙, [SW] 인용 안내, 팀원 확인 필요 항목 테이블',
     '전원'),
    ('해외선행연구.xlsx',
     '승원 정리 해외 선행연구 19편 (서지 확인 기준)',
     '영서·유나·승원'),
]
for ri, (a, b, c_) in enumerate(files):
    for ci, val in enumerate([a, b, c_]):
        cell = tbl3.rows[ri + 1].cells[ci]
        r = cell.paragraphs[0].add_run(val)
        r.font.name = FONT
        r.font.size = Pt(9)

doc.add_paragraph().paragraph_format.space_after = Pt(8)

# footer note
callout(
    '투고 여부 결정: 홍교수님 월요일(6/16) 피드백이 너무 절망적이지 않으면 한국지도학회지에 우선 투고 예정. '
    '1차 리비전 수정은 각 파트 담당자가 진행, 논문 등록·투고 심사 등 전반적인 행정 절차는 진이 담당합니다.',
    'EBEBEB', RGBColor(0x33, 0x33, 0x33), size=9
)

out = '/Users/jin/홍교수님 수업/RSG_final/공동작업_업무지시서.docx'
doc.save(out)
print(f'완료: {out}')
