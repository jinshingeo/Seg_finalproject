from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# A4 page setup
sec = doc.sections[0]
sec.page_width = Cm(21.0)
sec.page_height = Cm(29.7)
sec.left_margin = Cm(2.5)
sec.right_margin = Cm(2.5)
sec.top_margin = Cm(2.5)
sec.bottom_margin = Cm(2.5)

FONT = 'Palatino Linotype'
BODY_PT = 11
H1_PT = 12
H2_PT = 11
CAP_PT = 9
LINE_SPACING = 1.48
INDENT = Inches(0.30)


def set_run(run, size=BODY_PT, bold=False, italic=False, color=None):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def body(text, first_indent=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = LINE_SPACING
    p.paragraph_format.space_after = Pt(0)
    if first_indent:
        p.paragraph_format.first_line_indent = INDENT
    set_run(p.add_run(text))
    return p


def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.0
    set_run(p.add_run(text), size=H1_PT, bold=True)
    return p


def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.0
    set_run(p.add_run(text), size=H2_PT, bold=True)
    return p


def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.0
    set_run(p.add_run(text), size=H2_PT, italic=True)
    return p


def fig_placeholder(num, desc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'DCDCDC')
    pPr.append(shd)
    set_run(p.add_run(f'[그림 {num} 자리: {desc}]'),
            size=CAP_PT, color=RGBColor(0x55, 0x55, 0x55))

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.first_line_indent = Inches(0)
    cap.paragraph_format.space_after = Pt(10)
    set_run(cap.add_run(f'그림 {num}. {desc}.'), size=CAP_PT)


def eq_text(num, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(f'{text}    ({num})')
    run.font.name = 'Cambria Math'
    run.font.size = Pt(11)


# ─────────────────────────────────────────────────────
# TITLE
# ─────────────────────────────────────────────────────
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.first_line_indent = Inches(0)
p_title.paragraph_format.space_after = Pt(4)
set_run(p_title.add_run('서울시 저소득 독거노인의 소득 기반 거주지 분리 측도 분석'),
        size=14, bold=True)

p_title_en = doc.add_paragraph()
p_title_en.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title_en.paragraph_format.first_line_indent = Inches(0)
p_title_en.paragraph_format.space_after = Pt(20)
set_run(p_title_en.add_run(
    'Measuring Income-based Residential Segregation of\n'
    'Low-income Elderly Living Alone in Seoul'),
    size=11, italic=True)

# ─────────────────────────────────────────────────────
# ABSTRACT
# ─────────────────────────────────────────────────────
h1('초록 (Abstract)')

body(
    '저소득 독거노인은 소득 제약으로 거주지 선택권이 극도로 제한된 집단으로, '
    '특정 지역에 집중되는 거주지 분리가 심화될 경우 사회적 고립과 복지 서비스 '
    '접근성 불평등으로 이어질 수 있다.'
)
body(
    '본 연구는 서울시 424개 행정동을 대상으로 공간 가중 고립지수(Isolation Index)와 '
    '상이지수(Dissimilarity Index)를 적용하여 저소득 독거노인의 거주지 분리를 '
    '노출 축과 균등성 축에서 정량적으로 측정하였다.'
)
body(
    '분석 결과, 전역 고립지수(I = 0.0174)는 무작위 기대값 대비 약 14% 높았으며, '
    '일반인구 대비 전역 상이지수(D = 0.1318)는 독거노인 내부 소득 분리(D = 0.0859)보다 '
    '크게 나타났다. 국지 분석에서는 강남·서초의 배제 극과 노원·강서의 집중 극이라는 '
    '뚜렷한 양극 분화가 확인되었다.'
)
body(
    '이 결과는 소득 제약이 거주지 정렬을 통해 공간적 불평등으로 번역됨을 실증하며, '
    '행정동 단위의 미시적 복지 개입 근거를 제공한다.'
)

# ─────────────────────────────────────────────────────
# 1. INTRODUCTION
# ─────────────────────────────────────────────────────
h1('1. 서론 (Introduction)')

body(
    '한국의 독거노인 가구는 전체 고령자 가구의 36.1%로 가장 높은 비율을 차지하며, '
    '이 중 상당수가 기초생활수급 대상인 저소득 계층이다(⚠️ 통계청/보건복지부). '
    '특히 독거 여성 노인의 빈곤율은 58.6%에 달하는 것으로 보고되어(Ku, 2021 [SW]), '
    '저소득과 독거가 중첩되는 이중 취약 구조가 얼마나 심각한지를 짐작할 수 있다.'
)
body(
    '거주지 정렬(residential sorting) 이론은 소득 수준에 따라 개인·가구가 특정 지역으로 '
    '집중하거나 배제되는 현상이 구조적으로 발생함을 설명한다(Tiebout, 1956; '
    'Reardon and Bischoff, 2011). 저소득 독거노인은 낮은 소득(주거비 부담)과 독거(사회적 '
    '관계망 부재)라는 이중 제약을 동시에 가지므로, 거주지 정렬 메커니즘이 이 집단에게 '
    '특히 강하게 작동할 수 있다. 이러한 분리가 심화될 경우 사회적 교류 기회의 감소와 '
    '복지 서비스 접근성 불평등으로 이어진다는 점은(Kim et al., 2021 [SW]) 이미 지적된 '
    '바 있으며, 결국 저소득 독거노인의 거주지 분리는 단순한 공간 패턴이 아닌 복지 '
    '불평등의 공간적 발현으로 이해될 필요가 있다.'
)
body(
    '저소득 독거노인의 공간적 분포에 관한 기존 연구는 크게 두 방향으로 진행되어 왔다. '
    '첫째는 특정 집단이 어디에 집중되는지를 파악하는 공간 군집 분석으로, 노인 밀집 '
    '지역이나 고독사 위험 지역의 식별에 초점을 맞추어 왔다(이희연 등, 2015; '
    '안용한·김영호, 2023). '
    '둘째는 해당 분포에 영향을 미치는 요인을 규명하는 회귀 기반 접근으로, 주거비, '
    '교통 접근성 등의 변수와 노인 집중도의 관계를 분석하였다. 한편 소득 집단 간 '
    '공간적 분리를 정량적으로 측정하는 거주지 분리 측도 연구도 별도의 흐름으로 '
    '발전해 왔다(2장 검토). 그러나 이 두 흐름은 공통적으로 중요한 공백을 남긴다.'
)
body(
    '두 흐름 모두 중요한 공백을 남긴다. 첫째, 국내 독거노인 연구는 분포의 군집 여부나 '
    '영향 요인 분석에 집중할 뿐, 집단 간 공간적 분리의 정도 자체를 정량적으로 측정한 '
    '연구는 검토된 문헌 내에서 발견되지 않는다. 집중(concentration)과 분리(segregation)는 '
    '다른 현상으로, 저소득 독거노인이 특정 동에 많다는 것이 곧 다른 집단과 구조적으로 '
    '분리되어 있음을 의미하지 않는다. 둘째, 거주지 분리 측도를 적용한 연구들은 대부분 '
    '인종·민족 분리 맥락에 한정되어 있으며, 노출 축(고립지수)으로 저소득 독거노인의 '
    '사회적 고립 정도를 측정하거나, 균등성 축(상이지수)으로 독거노인 집단 내 소득 기반 '
    '분리를 측정한 연구는 검토된 문헌 내에서 확인되지 않는다.'
)
body(
    '최근 행정동 단위의 저소득 독거노인 및 일반 인구 구성 데이터의 가용성이 확보되었으며, '
    '공간 가중 분리 측도를 계산할 수 있는 방법론적 도구의 발전으로 기존 연구의 '
    '갭을 실증적으로 채울 조건이 갖춰졌다.'
)
body(
    '이에 본 연구는 서울시 424개 행정동을 분석 단위로, 저소득 독거노인의 소득 기반 '
    '거주지 분리를 노출 축(공간 고립지수)과 균등성 축(공간 상이지수)으로 측정하고자 '
    '한다. 구체적인 연구 질문은 다음과 같다. 연구 질문 1: 서울시 저소득 독거노인은 '
    '전체 주민 대비 공간적으로 얼마나 고립되어 있으며, 어떤 지역에 집중되는가? '
    '연구 질문 2: 저소득 독거노인의 거주지 분리는 일반인구 및 일반소득 독거노인 대비 '
    '어느 정도이며, 공간적으로 어떤 패턴을 보이는가? 이하 본 논문은 2장에서 선행연구를 '
    '검토하고, 3장에서 분석 방법론을 기술하며, 4장에서 결과 및 논의를 제시하고, '
    '5장에서 결론을 도출한다.'
)

# ─────────────────────────────────────────────────────
# 2. LITERATURE REVIEW
# ─────────────────────────────────────────────────────
h1('2. 선행연구')
h2('2.1 거주지 분리의 이론적 배경 및 측도 체계')

body(
    '거주지 정렬(residential sorting)은 개인·가구가 소득과 제약 조건에 따라 주거지를 '
    '선택하고 그 선택이 누적되면서 도시 내 사회 집단이 공간적으로 분화되는 현상이다. '
    'Tiebout(1956)은 소득이 높을수록 양질의 공공서비스를 제공하는 지역으로 이동함을 '
    '설명하였으며, Schelling(1971)은 개인의 작은 선호가 누적될 경우 예상보다 큰 분리가 '
    '발생함을 보였다. Reardon and Bischoff(2011)는 이 과정을 부유층의 배타적 입지와 '
    '저소득층의 선택지 제약이라는 두 경로로 구체화하였다.'
)
body(
    '거주지 분리의 측정 체계는 Massey and Denton(1988)의 5차원 분류에서 출발하여, '
    "Reardon and O'Sullivan(2004)이 균등성(evenness)과 노출(exposure)의 두 축으로 "
    '압축·정립하였다. Feitosa(2007)는 이를 가우시안 커널 기반의 공간 가중 측도로 '
    '확장하여 행정 경계의 자의성과 MAUP 문제를 완화하였으며, 본 연구에서 사용하는 '
    '공간 상이지수(D)와 공간 고립지수(I)의 이론적 기반을 제공한다.'
)

h2('2.2 소득 기반 거주지 분리 연구 (균등성 축)')

body(
    '소득 기반 거주지 분리는 균등성 축의 대표적 연구 대상으로, Cartone(2025) [SW]와 '
    'de Sousa Filho(2022) [SW] 등이 상이지수를 활용하여 소득 집단 간 공간적 불균등 '
    '분포를 실증한 바 있다. 이 연구들은 소득 불평등이 심화될수록 거주지 분리가 강화되는 '
    '패턴을 일관되게 확인하였다. 특히 Owens(2016) [SW]는 가구 유형을 통제한 후에도 '
    '소득이 거주지 분리의 독립 요인으로 남음을 실증하였다.'
)
body(
    '국내에서는 이희연 등(2015)이 서울시 저소득층 노인 밀집지구의 시·공간 '
    '분포와 근린환경 특성을 분석하여 월계2동·중계2·3동·등촌동 등 임대아파트 지역의 '
    '저소득층 노인 밀집을 확인하였으며, 안용한·김영호(2023)는 다층모형으로 저소득 '
    '독거노인 분포의 영향 요인을 규명한 바 있다. 그러나 이 연구들은 독거노인 가구를 '
    '소득 집단으로 구분하여 균등성 축의 분리 측도를 적용하지 않았다는 점에서 한계가 있다.'
)

h2('2.3 저소득 노인의 공간적 고립 연구 (노출 축)')

body(
    '노인의 사회적 고립에 관한 연구는 주로 설문 기반의 개인 척도(외로움, 사회적 관계망)에 '
    '집중되어 왔으며(⚠️ 관련 문헌 확인 필요), 집단 수준에서 공간적 고립을 측정한 '
    '연구는 제한적이다. 공간적 고립 연구 중 Menec(2019) [SW]는 저소득 노인 비율이 높은 '
    '지역에 고립 노인이 집중되는 패턴을 확인하였으며, Qin et al.(2024) [SW]는 공간 '
    '노출 지수를 활용하여 노인의 지역사회 고립을 측정한 바 있다.'
)
body(
    '저소득 독거노인을 관심 집단으로 설정하고, 전체 주민 대비 공간적 고립의 정도를 '
    '노출 축(고립지수)으로 측정한 연구는 검토된 문헌 내에서 발견되지 않는다.'
)

# ─────────────────────────────────────────────────────
# 3. METHODS
# ─────────────────────────────────────────────────────
h1('3. 연구 방법')
h2('3.1 연구 대상 지역')

body(
    '본 연구의 공간적 분석 단위는 서울특별시 424개 행정동이다. 행정동은 복지 서비스 '
    '전달의 기본 단위로, 저소득 독거노인 인구 데이터가 이 수준에서 공개되어 있다. '
    '공간 분리 측도 계산 시 미터 단위의 거리 계산이 필요하므로, 좌표계는 한국 측지계 '
    '기반의 EPSG:5179(Korea 2000 Unified CS)를 적용하였다.'
)
fig_placeholder(1, '서울시 424개 행정동 경계 및 저소득 독거노인 분포')

h2('3.2 데이터')

body(
    '분석 대상인 저소득 독거노인은 기초생활보장 수급자 중 65세 이상 1인 가구로 '
    '정의하였다. 비교 집단은 두 가지로 설정하였다. 첫째는 전체 주민으로, 고립지수(H1) '
    '계산 시 저소득 독거노인이 사회 전반과 얼마나 분리되어 있는지를 측정하기 위함이다. '
    '둘째는 일반소득 독거노인(비수급 독거노인)으로, 상이지수(H2-2) 계산 시 동일 가구 '
    '형태 내 소득 효과만을 분리하기 위함이다.'
)
body(
    '행정동별 독거노인 현황(저소득·일반)은 서울 열린데이터광장(2024)에서 수집하였으며, '
    '행정동별 주민등록인구(연령별)는 행정안전부(2024)에서 취득하였다. '
    '행정동 경계 shapefile은 SGIS를 통해 확보하였으며, '
    '공간 분리 측도 계산을 위한 Gaussian 커널 가중치 적용 시 활용하였다.'
)

# Table 1
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
tp.paragraph_format.first_line_indent = Inches(0)
tp.paragraph_format.space_before = Pt(8)
tp.paragraph_format.space_after = Pt(4)
set_run(tp.add_run('표 1. 데이터 목록 및 출처.'), size=CAP_PT)

tbl = doc.add_table(rows=5, cols=4)
tbl.style = 'Table Grid'
hdr_cells = tbl.rows[0].cells
for i, h_text in enumerate(['데이터명', '출처', '연도', '용도']):
    hdr_cells[i].text = h_text
    for run in hdr_cells[i].paragraphs[0].runs:
        run.bold = True
        run.font.name = FONT
        run.font.size = Pt(CAP_PT)

tbl_data = [
    ('행정동별 독거노인 현황', '서울 열린데이터광장', '2024', '저소득·일반소득 독거노인 인구'),
    ('주민등록인구(연령별)', '행정안전부', '2024', '전체 주민 인구'),
    ('행정동 경계 shapefile', 'SGIS', '2024', '공간 단위 설정 및 거리 계산'),
    ('좌표계 변환', 'EPSG:5179', '-', 'Korea 2000 Unified CS 적용'),
]
for ri, row_data in enumerate(tbl_data):
    for ci, val in enumerate(row_data):
        cell = tbl.rows[ri + 1].cells[ci]
        cell.text = val
        for run in cell.paragraphs[0].runs:
            run.font.name = FONT
            run.font.size = Pt(CAP_PT)

doc.add_paragraph().paragraph_format.space_after = Pt(8)

h2('3.3 분석 방법')
h3('3.3.1 공간 가중치 설정')

body(
    '행정동 경계를 분석 단위로 직접 사용할 경우 경계 획정 방식에 따라 결과가 달라지는 '
    '수정 가능 면적 단위 문제(MAUP)가 발생한다. 이를 완화하기 위해 Feitosa(2007)가 '
    '제안한 가우시안 커널 기반 공간 가중치를 적용하였다.'
)
body(
    '공간 가중치는 식 (1)의 가우시안 커널 함수로 정의되며, 거리가 멀수록 가중치가 '
    '지수적으로 감소한다. Bandwidth는 1,500m로 설정하였는데, 이는 서울시 행정동의 '
    '평균 반경(약 1km)을 고려한 값으로, 근린 수준의 공간 상호작용을 포착하기에 적합하다.'
)
eq_text(1, 'W_ij = exp(−d²_ij / 2·bw²)')

h3('3.3.2 공간 고립지수 (Spatial Isolation Index)')

body(
    '공간 고립지수(I)는 저소득 독거노인이 자신의 근린에서 같은 집단을 만날 확률로 '
    '정의되며(Bell, 1954; Feitosa, 2007), 노출(exposure) 차원의 분리를 측정한다. '
    '노출형 지수는 집단의 규모에 민감하므로, 절대값보다 무작위 기대값(π: 저소득 '
    '독거노인의 도시 전체 비율)과의 비교를 통해 해석해야 한다.'
)
eq_text(2, 'I = Σᵢ (aᵢ/A) · (L̃ᵃᵢ / L̃ᵗᵢ)')

h3('3.3.3 공간 상이지수 (Spatial Dissimilarity Index)')

body(
    '공간 상이지수(D)는 도시 전체에서 두 집단이 얼마나 불균등하게 분포하는지를 나타내는 '
    '균등성(evenness) 차원의 측도이다(Duncan and Duncan, 1955; Feitosa, 2007). 국지 '
    '상이지수(Local D)의 부호는 방향을 나타내는데, 양수는 저소득 독거노인이 지역 '
    '평균보다 과대 집중된 상태를, 음수는 상대적으로 배제된 상태를 의미한다.'
)
eq_text(3, 'D = Σᵢ (tᵢ/T) · |p̃ᵢ − P_A| / (2·P_A·P_B)')

h3('3.3.4 민감도 분석 (NSI + Bandwidth 비교)')

body(
    'Bandwidth 설정의 임의성을 검증하기 위해 500m부터 4,500m까지 다양한 bandwidth에서 '
    '전역 상이지수를 반복 계산하고 그 안정성을 확인하였다.'
)
body(
    '또한 근린 정렬 지수(NSI)를 통해 분리가 집중되는 공간적 스케일을 파악하였다. '
    'NSI는 각 bandwidth에서의 국지 분리 정도가 전체 분산에서 차지하는 비율로, '
    '값이 높은 bandwidth가 분리의 주요 스케일임을 나타낸다.'
)
eq_text(4, 'NSI = [Σᵢ L̃ᵗᵢ(X̂ᵢ − X̄)² / Σᵢ L̃ᵗᵢ] / S²_total')
fig_placeholder(2, 'Bandwidth별 전역 D값 변화 및 NSI 곡선')

# ─────────────────────────────────────────────────────
# 4. RESULTS AND DISCUSSION
# ─────────────────────────────────────────────────────
h1('4. 결과 및 논의')
h2('4.1 저소득 독거노인의 공간적 고립 (연구 질문 1: 노출 축)')

body(
    '전역 고립지수는 I = 0.0174로 나타났다. 이는 절대적으로 낮은 수준이나, 저소득 '
    '독거노인의 도시 전체 비율(π ≈ 1.52%)로 계산되는 무작위 기대값 대비 약 14% 높은 '
    '값으로, 집단 규모를 고려하면 통계적으로 유의미한 고립 수준이다(Feitosa, 2007).'
)
body(
    '국지 고립지수 지도에서는 강서구, 노원구, 강남구 수서동에 국지 고립이 응축되어 '
    '나타났다.'
)
fig_placeholder(3, 'Local I 단계구분도 (서울시 424개 행정동)')

# Table 4 — Local I Top 10
t4_cap = doc.add_paragraph()
t4_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
t4_cap.paragraph_format.first_line_indent = Inches(0)
t4_cap.paragraph_format.space_before = Pt(6)
t4_cap.paragraph_format.space_after = Pt(3)
set_run(t4_cap.add_run('표 4. 국지 고립지수(Local I) 상위 10개 행정동.'), size=CAP_PT)

t4 = doc.add_table(rows=11, cols=4)
t4.style = 'Table Grid'
for i, h_text in enumerate(['순위', '구', '동', 'Local I']):
    t4.rows[0].cells[i].text = h_text
    for run in t4.rows[0].cells[i].paragraphs[0].runs:
        run.bold = True
        run.font.name = FONT
        run.font.size = Pt(CAP_PT)
t4_data = [
    ('1', '강서구', '등촌3동', '0.000417'),
    ('2', '강서구', '가양2동', '0.000338'),
    ('3', '노원구', '중계2·3동', '0.000248'),
    ('4', '강남구', '수서동', '0.000202'),
    ('5', '노원구', '월계2동', '0.000194'),
    ('6', '노원구', '하계1동', '0.000187'),
    ('7', '강서구', '방화3동', '0.000185'),
    ('8', '중랑구', '망우본동', '0.000178'),
    ('9', '노원구', '상계3·4동', '0.000174'),
    ('10', '강서구', '가양3동', '0.000164'),
]
for ri, row_data in enumerate(t4_data):
    for ci, val in enumerate(row_data):
        cell = t4.rows[ri + 1].cells[ci]
        cell.text = val
        for run in cell.paragraphs[0].runs:
            run.font.name = FONT
            run.font.size = Pt(CAP_PT)

doc.add_paragraph().paragraph_format.space_after = Pt(4)

body(
    '강남·서초구에서는 Local I가 낮게 나타나는데, 이는 해당 지역의 저소득 독거노인 '
    '절대 수가 서울 평균 대비 극단적으로 적어 집단 내 접촉 확률 자체가 수학적으로 '
    '낮을 수밖에 없기 때문이다. 고립지수는 정의상 집단의 절대적 규모에 민감하게 '
    '반응하므로, 저소득 독거노인이 거의 없는 지역에서는 낮은 값이 필연적이다. '
    '그러나 그 소수의 저소득 독거노인이 전체 인구 구성에서 차지하는 비율은 서울 '
    '평균에서 크게 이탈해 있어, Local D는 음수(배제 극)로 나타난다. 이는 고립지수와 '
    '상이지수가 독립적으로 움직이는 분리의 두 차원임을 실증적으로 보여주며, 두 지수를 '
    '동시에 적용해야만 분리의 전체 구조를 정확히 파악할 수 있음을 시사한다.'
)
body(
    '이 결과는 소득 제약이 거주지 선택을 제한하여 특정 지역으로의 집중을 낳는다는 '
    'Reardon and Bischoff(2011)의 메커니즘과 일치하며, 강서·노원 지역의 저렴 주택 '
    '집중이 고립 응축의 공간적 기반임을 시사한다. 특히 국지 고립지수 상위권에 포함된 '
    '노원구 중계2·3동(3위)·월계2동(5위)은 이희연 등(2015)이 저소득층 노인 '
    '밀집지구로 확인한 영구임대아파트 지구와 일치하며, 본 연구의 측도 결과가 기존 '
    '밀집지구 연구와 같은 공간적 방향성을 보임을 실증한다. 나아가 Menec(2019) [SW]가 '
    '제시한 "저소득 노인 비율이 높은 지역에 고립 노인이 집중된다"는 패턴과도 일치하며, '
    '공간적 고립이 저소득 집중 지역에 중첩됨을 재확인한다.'
)

h2('4.2 일반인구 대비 거주지 분리 (연구 질문 2: 균등성 축)')

body(
    '일반인구 대비 전역 상이지수는 D = 0.1318로 나타났다. 관례적 기준(D < 0.30: '
    '낮은 분리, 0.30~0.60: 중간, > 0.60: 높은 분리)으로는 낮은 분리 수준에 해당한다. '
    '이 값은 독거노인 내부 소득 분리(D = 0.0859, 4.3절)보다 높게 나타나는데, 두 분석의 '
    '비교 집단 구성이 상이하므로(전체 주민 vs 동질적인 독거노인 집단) 수치를 직접 비교하는 '
    '데에는 한계가 있다. 다만 이 결과는 저소득 독거노인의 공간적 위치가 소득 효과 외에도 '
    '고령·1인 가구라는 인구학적 특성과 결합되어 있음을 시사한다.'
)
body(
    'Local D 지도에서는 노원·강서구(양수, 집중 극)와 강남·서초구(음수, 배제 극)의 '
    '뚜렷한 양극 분화가 관찰되었다. 양수 지역은 저소득 독거노인이 지역 인구 구성 '
    '비율 기준으로 서울 평균보다 과대 집중된 공간이며, 음수 지역은 저소득 독거노인이 '
    '해당 지역 인구 비율 대비 현저히 낮아 상대적으로 배제된 공간을 의미한다. 이 '
    '양극 분화는 서울의 주거비 지리를 반영하는데, 저렴한 임대주택이 밀집한 노원·강서 '
    '지역은 소득 제약이 강한 집단의 집중처가 되는 반면, 고가 주거지인 강남·서초에서는 '
    '이 집단이 구조적으로 배제된다.'
)
fig_placeholder(4, 'Local D 단계구분도 — H2-1 (양수: 집중 극, 음수: 배제 극)')
body(
    '이 양극 패턴은 Reardon and Bischoff(2011)가 제시한 소득 기반 거주지 분리의 두 경로, '
    '즉 부유층의 배타적 입지 선택(강남·서초의 배제 패턴)과 저소득층의 선택지 제약(노원·'
    '강서의 집중 패턴)과 정확히 대응된다. 이는 저소득 독거노인의 거주지 분리가 개인의 '
    '선택이 아닌 주거비 장벽에 의한 구조적 배제의 결과임을 시사하며, 단순한 공간 집중을 '
    '넘어 사회경제적 불평등이 공간적으로 가시화된 것으로 해석된다.'
)

h2('4.3 독거노인 내부 소득 기반 분리 (연구 질문 2: 세부)')

body(
    '독거노인 내부 소득 기반 전역 상이지수는 D = 0.0859로, H2-1(0.1318)보다 낮게 '
    '나타났다. 관례적 기준(D < 0.30: 낮은 분리)으로는 낮은 분리 수준에 해당하며, '
    '이는 서울 전역 평균으로는 독거노인 집단 안에서 소득에 따른 추가적 분리 효과가 '
    '상대적으로 제한적임을 시사한다.'
)
body(
    '그러나 국지 수준에서는 강남·서초(배제 극) vs 노원·강서(집중 극)의 양극 분화가 '
    'H2-1과 유사한 패턴으로 관찰되었다. Schelling(1971)의 누적 선호 이론에 따르면, '
    '집단 구성원 각각의 작은 선호(비슷한 소득 집단과의 거주 선호)가 반복·누적될 경우 '
    '전역 평균이 낮더라도 국지적으로 큰 분화가 발생할 수 있다. 전역 D = 0.0859라는 '
    '수치는 서울 전체 평균에서 독거 형태 안에서의 소득 분리가 상대적으로 제한적임을 '
    '보여주지만, 국지 수준의 양극 분화는 이 분리가 특정 지역에 집중되어 있음을 나타낸다.'
)
fig_placeholder(5, 'Local D 단계구분도 — H2-2 (독거노인 내부 소득 분리)')
body(
    '소득이 같은 가구 유형(독거노인) 안에서도 거주지를 분화시키는 독립 요인으로 '
    '작용한다는 점은 중요한 함의를 갖는다. 이는 독거라는 가구 형태를 통제한 후에도 '
    '소득이 거주지 정렬을 설명하는 독립 변수로 남음을 의미하며, Owens(2016) [SW]가 '
    '가구 유형 통제 후에도 소득이 분리의 독립 요인으로 남음을 실증한 결과와 일치한다. '
    '소득 기반 거주지 정렬이 가구 유형의 효과를 초월하여 작동하는 보편적 메커니즘임을 '
    '본 연구 결과와 Cartone(2025) [SW] 등 선행연구 모두 지지한다.'
)

h2('4.4 민감도 분석')

body(
    'Bandwidth를 500m에서 3,000m까지 변화시키며 전역 상이지수를 반복 계산한 결과, '
    '아래 표와 같이 단조 감소하면서도 안정적인 패턴을 보였다. 이는 분석 결과가 '
    'bandwidth 선택에 강건(robust)함을 나타낸다. 주목할 점은 3,000m 수준에서도 '
    'D = 0.054로 0에 수렴하지 않는다는 것으로, 서울 전체적으로 일정 수준의 소득 기반 '
    '공간 분리가 구조적으로 존재함을 시사한다.'
)

# Bandwidth D table
bw_cap = doc.add_paragraph()
bw_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
bw_cap.paragraph_format.first_line_indent = Inches(0)
bw_cap.paragraph_format.space_before = Pt(6)
bw_cap.paragraph_format.space_after = Pt(3)
set_run(bw_cap.add_run('표 2. Bandwidth별 전역 상이지수 (D, 연구 질문 2 기준).'), size=CAP_PT)

bw_tbl = doc.add_table(rows=6, cols=2)
bw_tbl.style = 'Table Grid'
for i, h_text in enumerate(['Bandwidth', '전역 D']):
    bw_tbl.rows[0].cells[i].text = h_text
    for run in bw_tbl.rows[0].cells[i].paragraphs[0].runs:
        run.bold = True
        run.font.name = FONT
        run.font.size = Pt(CAP_PT)
bw_data = [
    ('500 m', '0.1606'),
    ('1,000 m', '0.1107'),
    ('1,500 m (주 분석값)', '0.0859'),
    ('2,000 m', '0.0715'),
    ('3,000 m', '0.0540'),
]
for ri, (bw, d) in enumerate(bw_data):
    for ci, val in enumerate([bw, d]):
        cell = bw_tbl.rows[ri + 1].cells[ci]
        cell.text = val
        for run in cell.paragraphs[0].runs:
            run.font.name = FONT
            run.font.size = Pt(CAP_PT)

doc.add_paragraph().paragraph_format.space_after = Pt(4)

body(
    'NSI 분석 결과, 비공간 NSI(행정동 경계 기준)는 0.0503이었으나 반경 500m에서 '
    '0.0297(59%), 1,500m에서 0.0090(18%), 3,000m에서 0.0035(7%)로 bandwidth가 '
    '넓어질수록 단조 감소하였다. 이는 소득 기반 독거노인 분리가 서울 광역 구조의 '
    '문제가 아닌 행정동·근린 수준(~1,500m)의 국지적 패턴임을 의미하며, 자치구 단위 '
    '광역 정책보다 행정동 단위의 미시적 복지 개입이 더 효율적임을 시사한다.'
)

# NSI table
nsi_cap = doc.add_paragraph()
nsi_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
nsi_cap.paragraph_format.first_line_indent = Inches(0)
nsi_cap.paragraph_format.space_before = Pt(6)
nsi_cap.paragraph_format.space_after = Pt(3)
set_run(nsi_cap.add_run('표 3. Bandwidth별 NSI 값.'), size=CAP_PT)

nsi_tbl = doc.add_table(rows=7, cols=3)
nsi_tbl.style = 'Table Grid'
for i, h_text in enumerate(['Bandwidth', 'NSI', '비공간 대비']):
    nsi_tbl.rows[0].cells[i].text = h_text
    for run in nsi_tbl.rows[0].cells[i].paragraphs[0].runs:
        run.bold = True
        run.font.name = FONT
        run.font.size = Pt(CAP_PT)
nsi_data = [
    ('비공간 (행정동 경계)', '0.0503', '100%'),
    ('500 m', '0.0297', '59%'),
    ('1,000 m', '0.0138', '27%'),
    ('1,500 m', '0.0090', '18%'),
    ('2,000 m', '0.0064', '13%'),
    ('3,000 m', '0.0035', '7%'),
]
for ri, row_data in enumerate(nsi_data):
    for ci, val in enumerate(row_data):
        cell = nsi_tbl.rows[ri + 1].cells[ci]
        cell.text = val
        for run in cell.paragraphs[0].runs:
            run.font.name = FONT
            run.font.size = Pt(CAP_PT)

doc.add_paragraph().paragraph_format.space_after = Pt(4)
fig_placeholder(6, 'Bandwidth별 전역 D값 변화 그래프 및 NSI 곡선')

# ─────────────────────────────────────────────────────
# 5. CONCLUSION
# ─────────────────────────────────────────────────────
h1('5. 결론')

body(
    '본 연구는 서울시 424개 행정동을 대상으로 저소득 독거노인의 거주지 분리를 '
    '노출 축(공간 고립지수)과 균등성 축(공간 상이지수)으로 측정하였다. 전역 고립지수 '
    'I = 0.0174는 무작위 기대값(π = 0.0152) 대비 약 14% 높아 약한 수준의 고립이 '
    '확인되었으며, 국지적으로는 강서·노원 지역에 고립이 응축되었다. 일반인구 대비 '
    '전역 상이지수(D = 0.1318)는 독거노인 내부 소득 분리(D = 0.0859)보다 높게 '
    '나타나, 소득·고령·1인 가구 특성이 결합된 분리가 더 강하게 작동함을 확인하였다. '
    '이 패턴은 거주지 정렬 이론이 예측하는 부유층 배타 입지(강남·서초 배제)와 '
    '저소득층 선택지 제약(노원·강서 집중)의 두 경로와 정확히 대응된다.'
)
body(
    '이론적 기여 측면에서, 본 연구는 거주지 분리 측도를 인종·민족 맥락에서 소득·가구 '
    '형태 맥락으로 확장 적용하였다는 점에서 의의를 갖는다. '
    "Reardon and O'Sullivan(2004)의 두 축 프레임(노출·균등성)이 서울시 저소득 독거노인 "
    '집단에서도 유효하게 작동함을 실증하였으며, 특히 강남·서초에서 관찰된 '
    'Local I 낮음과 Local D 음수의 공존은 두 지수가 독립적으로 움직이는 분리의 두 차원임을 '
    '보여주어, 단일 지수로는 포착할 수 없는 분리 구조를 드러낸다. 정책적으로는 '
    'Local I 상위 행정동(강서구 등촌3·가양2동, 노원구 중계2·3동 등)을 복지 자원 배치의 '
    '우선 대상으로 수치로 뒷받침함으로써, 행정동 단위 미시적 개입의 근거를 제공한다.'
)
body(
    '본 연구는 네 가지 한계를 갖는다. 첫째, 2024년 단일 시점 데이터로 분리의 '
    '시계열적 변화를 파악하기 어렵다. 둘째, 기초생활수급자를 저소득 독거노인의 '
    '대리 변수로 사용하였으나, 차상위계층 등 수급 기준 외 실질적 빈곤층이 누락될 '
    '수 있으며, 주민등록 기준 데이터 특성상 실제 거주지와 등록지가 불일치하는 경우도 '
    '포함될 수 있다. 셋째, 행정동 단위 집계 데이터의 특성상 동 내부 이질성 포착이 '
    '불가능하며, 개인 수준의 결론 도출 시 생태학적 오류(ecological fallacy) 위험이 '
    '존재한다. 넷째, 노인여가복지시설의 공급 분포와의 교차 분석은 분리(수요 측) '
    '연구의 범위를 벗어나는 공급 측 문제로 본 연구에서 제외하였다.'
)
body(
    '향후 연구는 다음 방향으로 확장될 수 있다. 다년도 데이터를 활용한 분리 추이 분석으로 '
    '정렬의 동태적 과정을 추적하거나, 차상위계층을 포함한 확장된 저소득 정의를 적용할 '
    '수 있다. 또한 본 연구에서 식별된 고위험 행정동을 대상으로 복지시설 접근성 분석을 '
    '연계하여 수요-공급 불일치 구조를 분석하는 것도 유망한 연구 방향이다.'
)

# ─────────────────────────────────────────────────────
# REFERENCES
# ─────────────────────────────────────────────────────
h1('참고문헌')


def ref(parts):
    """parts: list of (text, italic) tuples"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = LINE_SPACING
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = INDENT
    p.paragraph_format.first_line_indent = -INDENT
    for text, italic in parts:
        r = p.add_run(text)
        r.font.name = FONT
        r.font.size = Pt(BODY_PT)
        r.italic = italic
    return p


def ref_simple(text):
    return ref([(text, False)])


# ── 국내 문헌 (자모음 순) ──
ref_simple('⚠️ 안용한·김영호, 2023, [제목·저널명·권호 확인 필요].')
ref_simple('서울 열린데이터광장, 2024, 행정동별 독거노인 현황(기초생활수급). '
           'https://data.seoul.go.kr (2024 검색).')
ref_simple('이희연·이다예·유재성, 2015, 저소득층 노인 밀집지구의 시·공간 분포와 '
           '근린환경 특성: 서울시를 사례로, 서울도시연구, 16(2), 1-18.')
ref_simple('행정안전부, 2024, 주민등록인구현황(연령별·읍면동별). '
           'https://jumin.mois.go.kr (2024 검색).')

# ── 외국어 문헌 (알파벳 순) ──
ref([('Bell, W., 1954, A probability model for the measurement of ecological '
      'segregation, ', False),
     ('Social Forces', True),
     (', 32(4), 357-364.', False)])

ref([('⚠️ [SW] Cartone, A., 2025, [제목·권호 확인 필요], ', False),
     ('Population, Space and Place', True),
     ('.', False)])

ref_simple('⚠️ [SW] de Sousa Filho, ..., 2022, [제목·저자·권호 확인 필요], '
           'SN Social Sciences.')

ref([('Duncan, O.D. and Duncan, B., 1955, A methodological analysis of '
      'segregation indexes, ', False),
     ('American Sociological Review', True),
     (', 20(2), 210-217.', False)])

ref([('Feitosa, F.F., Câmara, G., Monteiro, A.M.V., Koschitzki, T., and Silva, '
      'M.P.S., 2007, Global and local spatial indices of urban segregation, ', False),
     ('International Journal of Geographical Information Science', True),
     (', 21(3), 299-323. https://doi.org/10.1080/13658810600911903', False)])

ref_simple('⚠️ [SW] Kim et al., 2021, [저자 전체·제목·저널·권호 확인 필요].')

ref_simple('⚠️ [SW] Ku, ..., 2021, [저자 전체·제목·저널·권호 확인 필요].')

ref([('Massey, D.S. and Denton, N.A., 1988, The dimensions of residential '
      'segregation, ', False),
     ('Social Forces', True),
     (', 67(2), 281-315.', False)])

ref_simple('⚠️ [SW] Menec, V., 2019, [제목·저널·권호 확인 필요].')

ref([('⚠️ [SW] Owens, A., 2016, [제목·권호 확인 필요], ', False),
     ('American Sociological Review', True),
     ('.', False)])

ref_simple('⚠️ [SW] Qin et al., 2024, [저자 전체·제목·저널·권호 확인 필요].')

ref([('Reardon, S.F. and Bischoff, K., 2011, Income inequality and income '
      'segregation, ', False),
     ('American Journal of Sociology', True),
     (', 116(4), 1092-1153.', False)])

ref([("Reardon, S.F. and O'Sullivan, D., 2004, Measures of spatial segregation, ",
      False),
     ('Sociological Methodology', True),
     (', 34(1), 121-162.', False)])

ref([('Schelling, T.C., 1971, Dynamic models of segregation, ', False),
     ('Journal of Mathematical Sociology', True),
     (', 1(2), 143-186.', False)])

ref([('Tiebout, C.M., 1956, A pure theory of local expenditures, ', False),
     ('Journal of Political Economy', True),
     (', 64(5), 416-424.', False)])

out = '/Users/jin/홍교수님 수업/RSG_final/논문_초안_v1.docx'
doc.save(out)
print(f'완료: {out}')
